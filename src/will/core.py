"""
Schopenhauer Core - The WordDocument class.

This module provides the main WordDocument class for creating
and manipulating Word documents programmatically.

Example:
    >>> from will import WordDocument
    >>> doc = WordDocument()
    >>> doc.add_heading("My Report", level=1)
    >>> doc.add_paragraph("Introduction text here.")
    >>> doc.add_bullets(["Point 1", "Point 2", "Point 3"])
    >>> doc.save("report.docx")
"""

from io import BytesIO
from pathlib import Path
from typing import Any, Optional, Union

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt

from will.styles import (
    FONTS,
    MARGINS,
    PAGE_SIZES,
    SPACING,
    Colors,
    TableStyles,
)


class WordDocument:
    """
    A high-level interface for creating Word documents.

    The WordDocument class wraps python-docx to provide a simpler,
    more intuitive API for common document operations.

    Attributes:
        doc: The underlying python-docx Document object.

    Example:
        >>> doc = WordDocument()
        >>> doc.add_heading("Report Title", level=1)
        >>> doc.add_paragraph("This is the introduction.")
        >>> doc.save("report.docx")
    """

    def __init__(
        self,
        template: Optional[str] = None,
        title: Optional[str] = None,
        author: Optional[str] = None,
    ):
        """
        Initialize a new WordDocument.

        Args:
            template: Path to a .docx template file, or None for blank.
            title: Document title (stored in metadata).
            author: Document author (stored in metadata).
        """
        if template:
            template_path = Path(template)
            if not template_path.exists():
                raise FileNotFoundError(f"Template not found: {template}")
            self.doc = Document(template_path)
        else:
            self.doc = Document()

        # Set document properties
        if title:
            self.doc.core_properties.title = title
        if author:
            self.doc.core_properties.author = author

        # Apply default styling
        self._setup_default_styles()

    def _setup_default_styles(self):
        """Configure default styles for the document."""
        # Set default font for normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = FONTS.BODY.family
        font.size = FONTS.BODY.size_pt
        font.color.rgb = FONTS.BODY.color_rgb

        # Set default paragraph formatting
        para_format = style.paragraph_format
        para_format.space_after = SPACING.PARA_AFTER
        para_format.line_spacing = SPACING.LINE_DEFAULT

    # =========================================================================
    # HEADING METHODS
    # =========================================================================

    def add_heading(
        self,
        text: str,
        level: int = 1,
        color: Optional[str] = None,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    ):
        """
        Add a heading to the document.

        Args:
            text: The heading text.
            level: Heading level (0=Title, 1-9=Heading levels).
            color: Optional hex color override.
            alignment: Text alignment.

        Returns:
            The created paragraph object.
        """
        heading = self.doc.add_heading(text, level=level)

        # Apply styling based on level
        if level == 0:
            font_spec = FONTS.TITLE
        elif level == 1:
            font_spec = FONTS.HEADING_1
        elif level == 2:
            font_spec = FONTS.HEADING_2
        elif level == 3:
            font_spec = FONTS.HEADING_3
        else:
            font_spec = FONTS.HEADING_4

        for run in heading.runs:
            run.font.name = font_spec.family
            run.font.size = font_spec.size_pt
            run.font.bold = font_spec.bold
            if color:
                run.font.color.rgb = Colors.from_hex(color)
            else:
                run.font.color.rgb = font_spec.color_rgb

        heading.alignment = alignment
        return heading

    def add_title(
        self,
        text: str,
        subtitle: Optional[str] = None,
        centered: bool = True,
    ):
        """
        Add a document title (and optional subtitle).

        Args:
            text: The main title text.
            subtitle: Optional subtitle text.
            centered: Whether to center the title.

        Returns:
            Tuple of (title_paragraph, subtitle_paragraph or None)
        """
        alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT

        title = self.add_heading(text, level=0, alignment=alignment)

        subtitle_para = None
        if subtitle:
            subtitle_para = self.doc.add_paragraph()
            run = subtitle_para.add_run(subtitle)
            run.font.name = FONTS.SUBTITLE.family
            run.font.size = FONTS.SUBTITLE.size_pt
            run.font.italic = FONTS.SUBTITLE.italic
            run.font.color.rgb = FONTS.SUBTITLE.color_rgb
            subtitle_para.alignment = alignment

        return title, subtitle_para

    # =========================================================================
    # PARAGRAPH METHODS
    # =========================================================================

    def add_paragraph(
        self,
        text: str = "",
        style: Optional[str] = None,
        bold: bool = False,
        italic: bool = False,
        font_size: Optional[int] = None,
        font_name: Optional[str] = None,
        color: Optional[str] = None,
        alignment: Optional[WD_ALIGN_PARAGRAPH] = None,
        first_line_indent: Optional[float] = None,
        space_before: Optional[int] = None,
        space_after: Optional[int] = None,
    ):
        """
        Add a paragraph to the document.

        Args:
            text: The paragraph text.
            style: Built-in style name to apply.
            bold: Whether text should be bold.
            italic: Whether text should be italic.
            font_size: Font size in points.
            font_name: Font family name.
            color: Hex color for the text.
            alignment: Text alignment.
            first_line_indent: First line indent in inches.
            space_before: Space before paragraph in points.
            space_after: Space after paragraph in points.

        Returns:
            The created paragraph object.
        """
        para = self.doc.add_paragraph(style=style)

        if text:
            run = para.add_run(text)
            run.bold = bold
            run.italic = italic

            if font_size:
                run.font.size = Pt(font_size)
            if font_name:
                run.font.name = font_name
            if color:
                run.font.color.rgb = Colors.from_hex(color)

        if alignment:
            para.alignment = alignment

        para_format = para.paragraph_format
        if first_line_indent is not None:
            para_format.first_line_indent = Inches(first_line_indent)
        if space_before is not None:
            para_format.space_before = Pt(space_before)
        if space_after is not None:
            para_format.space_after = Pt(space_after)

        return para

    def add_text(
        self,
        paragraph,
        text: str,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        font_size: Optional[int] = None,
        font_name: Optional[str] = None,
        color: Optional[str] = None,
    ):
        """
        Add formatted text to an existing paragraph.

        Args:
            paragraph: The paragraph to add text to.
            text: The text to add.
            bold: Whether text should be bold.
            italic: Whether text should be italic.
            underline: Whether text should be underlined.
            font_size: Font size in points.
            font_name: Font family name.
            color: Hex color for the text.

        Returns:
            The created run object.
        """
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline

        if font_size:
            run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name
        if color:
            run.font.color.rgb = Colors.from_hex(color)

        return run

    # =========================================================================
    # LIST METHODS
    # =========================================================================

    def add_bullets(
        self,
        items: list[str],
        level: int = 0,
        style: str = "List Bullet",
    ):
        """
        Add a bulleted list to the document.

        Args:
            items: List of strings for bullet points.
            level: Indentation level (0-8).
            style: The bullet style to use.

        Returns:
            List of created paragraph objects.
        """
        paragraphs = []
        for item in items:
            para = self.doc.add_paragraph(item, style=style)
            # Set indentation level
            if level > 0:
                para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            paragraphs.append(para)
        return paragraphs

    def add_numbered_list(
        self,
        items: list[str],
        level: int = 0,
        style: str = "List Number",
    ):
        """
        Add a numbered list to the document.

        Args:
            items: List of strings for numbered items.
            level: Indentation level (0-8).
            style: The numbering style to use.

        Returns:
            List of created paragraph objects.
        """
        paragraphs = []
        for item in items:
            para = self.doc.add_paragraph(item, style=style)
            if level > 0:
                para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            paragraphs.append(para)
        return paragraphs

    # =========================================================================
    # TABLE METHODS
    # =========================================================================

    def add_table(
        self,
        data: list[list[Any]],
        headers: Optional[list[str]] = None,
        style: str = "Table Grid",
        header_color: Optional[str] = None,
        alternating_rows: bool = True,
        column_widths: Optional[list[float]] = None,
    ):
        """
        Add a table to the document.

        Args:
            data: 2D list of cell values.
            headers: Optional list of header strings.
            style: Table style name.
            header_color: Hex color for header background.
            alternating_rows: Whether to shade alternate rows.
            column_widths: List of column widths in inches.

        Returns:
            The created table object.
        """
        num_cols = len(headers) if headers else len(data[0]) if data else 0
        num_rows = len(data) + (1 if headers else 0)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = style
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths if provided
        if column_widths:
            for i, width in enumerate(column_widths):
                if i < len(table.columns):
                    for cell in table.columns[i].cells:
                        cell.width = Inches(width)

        row_offset = 0

        # Add headers
        if headers:
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = str(header)
                # Apply header styling
                color = Colors.from_hex(header_color) if header_color else Colors.PRIMARY
                TableStyles.apply_header_style(cell, color)
            row_offset = 1

        # Add data rows
        for row_idx, row_data in enumerate(data):
            table_row = table.rows[row_idx + row_offset]
            for col_idx, cell_value in enumerate(row_data):
                cell = table_row.cells[col_idx]
                cell.text = str(cell_value)
                TableStyles.apply_body_style(cell)

        # Apply alternating row colors
        if alternating_rows:
            TableStyles.apply_alternating_rows(table)

        return table

    # =========================================================================
    # IMAGE METHODS
    # =========================================================================

    def add_image(
        self,
        image_path: str,
        width: Optional[float] = None,
        height: Optional[float] = None,
        caption: Optional[str] = None,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    ):
        """
        Add an image to the document.

        Args:
            image_path: Path to the image file.
            width: Image width in inches (maintains aspect ratio if height not set).
            height: Image height in inches.
            caption: Optional caption text below the image.
            alignment: Alignment of the image paragraph.

        Returns:
            Tuple of (paragraph, inline_shape, caption_paragraph or None)
        """
        para = self.doc.add_paragraph()
        para.alignment = alignment

        # Determine dimensions
        width_inches = Inches(width) if width else None
        height_inches = Inches(height) if height else None

        run = para.add_run()
        picture = run.add_picture(image_path, width=width_inches, height=height_inches)

        caption_para = None
        if caption:
            caption_para = self.doc.add_paragraph()
            run = caption_para.add_run(caption)
            run.font.name = FONTS.CAPTION.family
            run.font.size = FONTS.CAPTION.size_pt
            run.font.italic = FONTS.CAPTION.italic
            run.font.color.rgb = FONTS.CAPTION.color_rgb
            caption_para.alignment = alignment

        return para, picture, caption_para

    # =========================================================================
    # SPECIAL CONTENT METHODS
    # =========================================================================

    def add_quote(
        self,
        text: str,
        author: Optional[str] = None,
        style: str = "Quote",
    ):
        """
        Add a blockquote to the document.

        Args:
            text: The quote text.
            author: Optional attribution.
            style: The quote style to use.

        Returns:
            List of created paragraphs.
        """
        paragraphs = []

        # Add the quote
        quote_para = self.doc.add_paragraph(style=style)
        run = quote_para.add_run(f'"{text}"')
        run.font.name = FONTS.QUOTE.family
        run.font.size = FONTS.QUOTE.size_pt
        run.font.italic = True
        run.font.color.rgb = FONTS.QUOTE.color_rgb
        quote_para.paragraph_format.left_indent = Inches(0.5)
        quote_para.paragraph_format.right_indent = Inches(0.5)
        paragraphs.append(quote_para)

        # Add attribution
        if author:
            attr_para = self.doc.add_paragraph()
            run = attr_para.add_run(f"— {author}")
            run.font.name = FONTS.CAPTION.family
            run.font.size = FONTS.CAPTION.size_pt
            run.font.italic = True
            run.font.color.rgb = FONTS.CAPTION.color_rgb
            attr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            attr_para.paragraph_format.right_indent = Inches(0.5)
            paragraphs.append(attr_para)

        return paragraphs

    def add_code_block(
        self,
        code: str,
        language: Optional[str] = None,
    ):
        """
        Add a code block to the document.

        Args:
            code: The code text.
            language: Optional language identifier (for documentation).

        Returns:
            The created paragraph.
        """
        para = self.doc.add_paragraph()
        run = para.add_run(code)
        run.font.name = FONTS.CODE.family
        run.font.size = FONTS.CODE.size_pt
        run.font.color.rgb = FONTS.CODE.color_rgb

        # Add light gray background
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.left_indent = Inches(0.25)

        return para

    def add_horizontal_line(self):
        """Add a horizontal line (border) to the document."""
        para = self.doc.add_paragraph()
        # Add bottom border to simulate horizontal line
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="{Colors.GRAY_400.value}"/>'
            f'</w:pBdr>'
        )
        para._p.get_or_add_pPr().append(pBdr)
        return para

    def add_page_break(self):
        """Add a page break to the document."""
        self.doc.add_page_break()

    def add_section_break(self, start_type=None):
        """Add a section break to the document."""
        self.doc.add_section(start_type)

    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================

    def add_table_of_contents(
        self,
        title: str = "Table of Contents",
        levels: int = 3,
    ):
        """
        Add a table of contents field to the document.

        Note: The TOC will need to be updated when opened in Word.

        Args:
            title: Title for the TOC section.
            levels: Number of heading levels to include.

        Returns:
            The TOC paragraph.
        """
        # Add TOC title
        self.add_heading(title, level=1)

        # Add TOC field
        para = self.doc.add_paragraph()

        # Create TOC field code
        run = para.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)

        run2 = para.add_run()
        instrText = parse_xml(
            f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-{levels}" \\h \\z \\u </w:instrText>'
        )
        run2._r.append(instrText)

        run3 = para.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run3._r.append(fldChar2)

        run4 = para.add_run("Right-click and select 'Update Field' to generate TOC")
        run4.font.color.rgb = Colors.GRAY_500.rgb
        run4.italic = True

        run5 = para.add_run()
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run5._r.append(fldChar3)

        # Add page break after TOC
        self.add_page_break()

        return para

    # =========================================================================
    # HEADER/FOOTER METHODS
    # =========================================================================

    def set_header(
        self,
        text: str,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    ):
        """
        Set the document header text.

        Args:
            text: Header text.
            alignment: Text alignment.

        Returns:
            The header paragraph.
        """
        section = self.doc.sections[0]
        header = section.header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.clear()

        run = para.add_run(text)
        run.font.name = FONTS.HEADER.family
        run.font.size = FONTS.HEADER.size_pt
        run.font.color.rgb = FONTS.HEADER.color_rgb
        para.alignment = alignment

        return para

    def set_footer(
        self,
        text: str = "",
        include_page_numbers: bool = True,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    ):
        """
        Set the document footer.

        Args:
            text: Footer text.
            include_page_numbers: Whether to include page numbers.
            alignment: Text alignment.

        Returns:
            The footer paragraph.
        """
        section = self.doc.sections[0]
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()

        if text:
            run = para.add_run(text)
            run.font.name = FONTS.FOOTER.family
            run.font.size = FONTS.FOOTER.size_pt
            run.font.color.rgb = FONTS.FOOTER.color_rgb

        if include_page_numbers:
            if text:
                para.add_run(" | ")

            # Add page number field
            run = para.add_run()
            fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
            run._r.append(fldChar1)

            run2 = para.add_run()
            instrText = parse_xml(
                f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'
            )
            run2._r.append(instrText)

            run3 = para.add_run()
            fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
            run3._r.append(fldChar2)

        para.alignment = alignment
        return para

    # =========================================================================
    # PAGE SETUP METHODS
    # =========================================================================

    def set_margins(
        self,
        top: Optional[float] = None,
        bottom: Optional[float] = None,
        left: Optional[float] = None,
        right: Optional[float] = None,
        preset: str = None,
    ):
        """
        Set page margins.

        Args:
            top: Top margin in inches.
            bottom: Bottom margin in inches.
            left: Left margin in inches.
            right: Right margin in inches.
            preset: Use a preset ("normal", "narrow", "moderate", "wide").
        """
        section = self.doc.sections[0]

        if preset:
            presets = {
                "normal": MARGINS.NORMAL,
                "narrow": MARGINS.NARROW,
                "moderate": MARGINS.MODERATE,
                "wide": MARGINS.WIDE,
            }
            margins = presets.get(preset.lower(), MARGINS.NORMAL)
            section.top_margin = margins["top"]
            section.bottom_margin = margins["bottom"]
            section.left_margin = margins["left"]
            section.right_margin = margins["right"]
        else:
            if top is not None:
                section.top_margin = Inches(top)
            if bottom is not None:
                section.bottom_margin = Inches(bottom)
            if left is not None:
                section.left_margin = Inches(left)
            if right is not None:
                section.right_margin = Inches(right)

    def set_page_size(
        self,
        width: Optional[float] = None,
        height: Optional[float] = None,
        preset: str = None,
        orientation: str = "portrait",
    ):
        """
        Set page size.

        Args:
            width: Page width in inches.
            height: Page height in inches.
            preset: Use a preset ("letter", "legal", "a4", "a5").
            orientation: "portrait" or "landscape".
        """
        section = self.doc.sections[0]

        if preset:
            presets = {
                "letter": PAGE_SIZES.LETTER,
                "legal": PAGE_SIZES.LEGAL,
                "a4": PAGE_SIZES.A4,
                "a5": PAGE_SIZES.A5,
            }
            size = presets.get(preset.lower(), PAGE_SIZES.LETTER)
            section.page_width = size["width"]
            section.page_height = size["height"]
        else:
            if width is not None:
                section.page_width = Inches(width)
            if height is not None:
                section.page_height = Inches(height)

        if orientation.lower() == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            # Swap width and height for landscape
            section.page_width, section.page_height = section.page_height, section.page_width

    # =========================================================================
    # TEMPLATE/PLACEHOLDER METHODS
    # =========================================================================

    def replace_placeholders(self, replacements: dict[str, str]):
        """
        Replace {{PLACEHOLDER}} tokens throughout the document.

        Args:
            replacements: Dictionary of placeholder names to values.
                         Example: {"NAME": "John", "DATE": "2024-01-15"}

        Returns:
            Number of replacements made.
        """
        count = 0

        # Replace in paragraphs
        for para in self.doc.paragraphs:
            for key, value in replacements.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in para.text:
                    # Need to handle runs carefully
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
                            count += 1

        # Replace in tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in replacements.items():
                            placeholder = f"{{{{{key}}}}}"
                            for run in para.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
                                    count += 1

        # Replace in headers/footers
        for section in self.doc.sections:
            for para in section.header.paragraphs:
                for key, value in replacements.items():
                    placeholder = f"{{{{{key}}}}}"
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
                            count += 1

            for para in section.footer.paragraphs:
                for key, value in replacements.items():
                    placeholder = f"{{{{{key}}}}}"
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))
                            count += 1

        return count

    def get_placeholders(self) -> list[str]:
        """
        Find all {{PLACEHOLDER}} tokens in the document.

        Returns:
            List of placeholder names found.
        """
        import re
        placeholders = set()
        pattern = r'\{\{([^}]+)\}\}'

        # Check paragraphs
        for para in self.doc.paragraphs:
            matches = re.findall(pattern, para.text)
            placeholders.update(matches)

        # Check tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = re.findall(pattern, para.text)
                        placeholders.update(matches)

        # Check headers/footers
        for section in self.doc.sections:
            for para in section.header.paragraphs:
                matches = re.findall(pattern, para.text)
                placeholders.update(matches)
            for para in section.footer.paragraphs:
                matches = re.findall(pattern, para.text)
                placeholders.update(matches)

        return sorted(placeholders)

    # =========================================================================
    # DOCUMENT INFO METHODS
    # =========================================================================

    def get_info(self) -> dict[str, Any]:
        """
        Get document metadata and statistics.

        Returns:
            Dictionary containing document information.
        """
        info = {
            "title": self.doc.core_properties.title or "",
            "author": self.doc.core_properties.author or "",
            "created": str(self.doc.core_properties.created) if self.doc.core_properties.created else "",
            "modified": str(self.doc.core_properties.modified) if self.doc.core_properties.modified else "",
            "paragraphs": len(self.doc.paragraphs),
            "tables": len(self.doc.tables),
            "sections": len(self.doc.sections),
            "styles": [s.name for s in self.doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH],
            "placeholders": self.get_placeholders(),
        }

        # Word count
        word_count = 0
        for para in self.doc.paragraphs:
            word_count += len(para.text.split())
        info["word_count"] = word_count

        return info

    # =========================================================================
    # SAVE METHODS
    # =========================================================================

    def save(self, path: Union[str, Path]):
        """
        Save the document to a file.

        Args:
            path: Output file path.
        """
        self.doc.save(str(path))

    def to_bytes(self) -> bytes:
        """
        Get the document as bytes.

        Returns:
            Document content as bytes.
        """
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    # =========================================================================
    # CLASS METHODS
    # =========================================================================

    @classmethod
    def from_yaml(cls, yaml_path: str, output_path: Optional[str] = None) -> "WordDocument":
        """
        Create a document from a YAML specification.

        Args:
            yaml_path: Path to YAML file.
            output_path: Optional path to save the document.

        Returns:
            The created WordDocument.
        """
        import yaml

        with open(yaml_path, encoding='utf-8') as f:
            spec = yaml.safe_load(f)

        doc = cls.from_spec(spec)

        if output_path:
            doc.save(output_path)

        return doc

    @classmethod
    def from_json(cls, json_path: str, output_path: Optional[str] = None) -> "WordDocument":
        """
        Create a document from a JSON specification.

        Args:
            json_path: Path to JSON file.
            output_path: Optional path to save the document.

        Returns:
            The created WordDocument.
        """
        import json

        with open(json_path, encoding='utf-8') as f:
            spec = json.load(f)

        doc = cls.from_spec(spec)

        if output_path:
            doc.save(output_path)

        return doc

    @classmethod
    def from_spec(cls, spec: dict[str, Any]) -> "WordDocument":
        """
        Create a document from a specification dictionary.

        Args:
            spec: Document specification dictionary.

        Returns:
            The created WordDocument.
        """
        # Get document-level settings
        template = spec.get("template")
        title = spec.get("title", "")
        subtitle = spec.get("subtitle")
        author = spec.get("author")
        page_size = spec.get("page_size", "letter")
        margins = spec.get("margins", "normal")
        header = spec.get("header")
        footer = spec.get("footer")
        include_toc = spec.get("table_of_contents", False)

        # Create document
        doc = cls(template=template, title=title, author=author)

        # Page setup
        doc.set_page_size(preset=page_size)
        doc.set_margins(preset=margins)

        # Header/footer
        if header:
            doc.set_header(header)
        if footer:
            doc.set_footer(footer)

        # Title page
        if title:
            doc.add_title(title, subtitle)
            if spec.get("title_page_break", True):
                doc.add_page_break()

        # Table of contents
        if include_toc:
            doc.add_table_of_contents()

        # Process sections/content
        sections = spec.get("sections", [])
        for section in sections:
            doc._process_section(section)

        return doc

    def _process_section(self, section: dict[str, Any]):
        """Process a section from the spec."""
        section_type = section.get("type", "content")

        if section_type == "heading":
            level = section.get("level", 1)
            self.add_heading(section.get("title", ""), level=level)

        elif section_type == "section":
            # Section header with optional page break
            if section.get("page_break", False):
                self.add_page_break()
            self.add_heading(section.get("title", ""), level=1)
            if section.get("subtitle"):
                self.add_paragraph(section["subtitle"], italic=True)

        elif section_type == "content":
            if section.get("title"):
                level = section.get("level", 2)
                self.add_heading(section["title"], level=level)

            if section.get("text"):
                self.add_paragraph(section["text"])

            if section.get("bullets"):
                self.add_bullets(section["bullets"])

            if section.get("numbered"):
                self.add_numbered_list(section["numbered"])

        elif section_type == "table":
            if section.get("title"):
                self.add_heading(section["title"], level=2)
            self.add_table(
                data=section.get("data", []),
                headers=section.get("headers"),
                column_widths=section.get("column_widths"),
            )

        elif section_type == "image":
            self.add_image(
                image_path=section.get("path", section.get("image", "")),
                width=section.get("width"),
                height=section.get("height"),
                caption=section.get("caption", section.get("title")),
            )

        elif section_type == "quote":
            self.add_quote(
                text=section.get("text", ""),
                author=section.get("author"),
            )

        elif section_type == "code":
            if section.get("title"):
                self.add_heading(section["title"], level=3)
            self.add_code_block(
                code=section.get("code", section.get("text", "")),
                language=section.get("language"),
            )

        elif section_type == "page_break":
            self.add_page_break()

        elif section_type == "horizontal_line":
            self.add_horizontal_line()
