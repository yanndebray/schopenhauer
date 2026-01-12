"""
Schopenhauer Document Builder - Fluent interface for document construction.

This module provides the DocumentBuilder class, which offers a fluent
(method chaining) interface for building Word documents incrementally.

Example:
    >>> from will import WordDocument, DocumentBuilder
    >>> doc = WordDocument()
    >>> builder = DocumentBuilder(doc)
    >>> (builder
    ...     .set_title("My Report")
    ...     .add_heading("Introduction", level=1)
    ...     .add_paragraph("This is the intro.")
    ...     .add_bullets(["Point 1", "Point 2"])
    ...     .add_table([["A", "B"], ["C", "D"]], headers=["Col1", "Col2"])
    ...     .save("report.docx"))
"""

from pathlib import Path
from typing import Any, Optional, Union

from docx.enum.text import WD_ALIGN_PARAGRAPH

from will.core import WordDocument


class DocumentBuilder:
    """
    Fluent builder for creating Word documents.

    The DocumentBuilder provides a chainable interface for document creation,
    making it easy to construct complex documents with readable code.

    Example:
        >>> builder = DocumentBuilder()
        >>> builder.set_title("Report").add_paragraph("Hello").save("out.docx")
    """

    def __init__(self, document: Optional[WordDocument] = None):
        """
        Initialize the DocumentBuilder.

        Args:
            document: Optional existing WordDocument to build on.
                     If None, creates a new document.
        """
        self._doc = document or WordDocument()
        self._current_paragraph = None
        self._current_table = None

    @property
    def document(self) -> WordDocument:
        """Get the underlying WordDocument."""
        return self._doc

    # =========================================================================
    # DOCUMENT SETUP METHODS
    # =========================================================================

    def set_title(
        self,
        title: str,
        subtitle: Optional[str] = None,
        centered: bool = True,
    ) -> "DocumentBuilder":
        """
        Set the document title.

        Args:
            title: Main title text.
            subtitle: Optional subtitle.
            centered: Whether to center the title.

        Returns:
            Self for method chaining.
        """
        self._doc.add_title(title, subtitle, centered)
        return self

    def set_author(self, author: str) -> "DocumentBuilder":
        """
        Set the document author.

        Args:
            author: Author name.

        Returns:
            Self for method chaining.
        """
        self._doc.doc.core_properties.author = author
        return self

    def set_header(
        self,
        text: str,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    ) -> "DocumentBuilder":
        """
        Set the document header.

        Args:
            text: Header text.
            alignment: Text alignment.

        Returns:
            Self for method chaining.
        """
        self._doc.set_header(text, alignment)
        return self

    def set_footer(
        self,
        text: str = "",
        include_page_numbers: bool = True,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    ) -> "DocumentBuilder":
        """
        Set the document footer.

        Args:
            text: Footer text.
            include_page_numbers: Whether to include page numbers.
            alignment: Text alignment.

        Returns:
            Self for method chaining.
        """
        self._doc.set_footer(text, include_page_numbers, alignment)
        return self

    def set_margins(
        self,
        top: Optional[float] = None,
        bottom: Optional[float] = None,
        left: Optional[float] = None,
        right: Optional[float] = None,
        preset: Optional[str] = None,
    ) -> "DocumentBuilder":
        """
        Set page margins.

        Args:
            top: Top margin in inches.
            bottom: Bottom margin in inches.
            left: Left margin in inches.
            right: Right margin in inches.
            preset: Preset name ("normal", "narrow", "moderate", "wide").

        Returns:
            Self for method chaining.
        """
        self._doc.set_margins(top, bottom, left, right, preset)
        return self

    def set_page_size(
        self,
        width: Optional[float] = None,
        height: Optional[float] = None,
        preset: Optional[str] = None,
        orientation: str = "portrait",
    ) -> "DocumentBuilder":
        """
        Set page size.

        Args:
            width: Page width in inches.
            height: Page height in inches.
            preset: Preset name ("letter", "legal", "a4", "a5").
            orientation: "portrait" or "landscape".

        Returns:
            Self for method chaining.
        """
        self._doc.set_page_size(width, height, preset, orientation)
        return self

    # =========================================================================
    # HEADING METHODS
    # =========================================================================

    def add_heading(
        self,
        text: str,
        level: int = 1,
        color: Optional[str] = None,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    ) -> "DocumentBuilder":
        """
        Add a heading.

        Args:
            text: Heading text.
            level: Heading level (0=Title, 1-9=Heading levels).
            color: Optional hex color override.
            alignment: Text alignment.

        Returns:
            Self for method chaining.
        """
        self._current_paragraph = self._doc.add_heading(text, level, color, alignment)
        return self

    def add_section(
        self,
        title: str,
        subtitle: Optional[str] = None,
        page_break: bool = False,
    ) -> "DocumentBuilder":
        """
        Add a new section with optional page break.

        Args:
            title: Section title.
            subtitle: Optional subtitle.
            page_break: Whether to add a page break before.

        Returns:
            Self for method chaining.
        """
        if page_break:
            self._doc.add_page_break()

        self._doc.add_heading(title, level=1)

        if subtitle:
            self._doc.add_paragraph(subtitle, italic=True)

        return self

    # =========================================================================
    # CONTENT METHODS
    # =========================================================================

    def add_paragraph(
        self,
        text: str = "",
        bold: bool = False,
        italic: bool = False,
        font_size: Optional[int] = None,
        color: Optional[str] = None,
        alignment: Optional[WD_ALIGN_PARAGRAPH] = None,
    ) -> "DocumentBuilder":
        """
        Add a paragraph.

        Args:
            text: Paragraph text.
            bold: Whether text should be bold.
            italic: Whether text should be italic.
            font_size: Font size in points.
            color: Hex color for the text.
            alignment: Text alignment.

        Returns:
            Self for method chaining.
        """
        self._current_paragraph = self._doc.add_paragraph(
            text=text,
            bold=bold,
            italic=italic,
            font_size=font_size,
            color=color,
            alignment=alignment,
        )
        return self

    def add_text(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        color: Optional[str] = None,
    ) -> "DocumentBuilder":
        """
        Add text to the current paragraph.

        Args:
            text: Text to add.
            bold: Whether text should be bold.
            italic: Whether text should be italic.
            underline: Whether text should be underlined.
            color: Hex color for the text.

        Returns:
            Self for method chaining.
        """
        if not self._current_paragraph:
            self._current_paragraph = self._doc.add_paragraph()

        self._doc.add_text(
            self._current_paragraph,
            text,
            bold=bold,
            italic=italic,
            underline=underline,
            color=color,
        )
        return self

    def add_bullets(
        self,
        items: list[str],
        level: int = 0,
    ) -> "DocumentBuilder":
        """
        Add a bulleted list.

        Args:
            items: List of bullet point strings.
            level: Indentation level.

        Returns:
            Self for method chaining.
        """
        self._doc.add_bullets(items, level)
        return self

    def add_numbered_list(
        self,
        items: list[str],
        level: int = 0,
    ) -> "DocumentBuilder":
        """
        Add a numbered list.

        Args:
            items: List of numbered item strings.
            level: Indentation level.

        Returns:
            Self for method chaining.
        """
        self._doc.add_numbered_list(items, level)
        return self

    # =========================================================================
    # TABLE METHODS
    # =========================================================================

    def add_table(
        self,
        data: list[list[Any]],
        headers: Optional[list[str]] = None,
        header_color: Optional[str] = None,
        alternating_rows: bool = True,
        column_widths: Optional[list[float]] = None,
    ) -> "DocumentBuilder":
        """
        Add a table.

        Args:
            data: 2D list of cell values.
            headers: Optional list of header strings.
            header_color: Hex color for header background.
            alternating_rows: Whether to shade alternate rows.
            column_widths: List of column widths in inches.

        Returns:
            Self for method chaining.
        """
        self._current_table = self._doc.add_table(
            data=data,
            headers=headers,
            header_color=header_color,
            alternating_rows=alternating_rows,
            column_widths=column_widths,
        )
        return self

    def add_key_value_table(
        self,
        data: dict[str, Any],
        key_header: str = "Property",
        value_header: str = "Value",
    ) -> "DocumentBuilder":
        """
        Add a key-value table from a dictionary.

        Args:
            data: Dictionary of key-value pairs.
            key_header: Header for the key column.
            value_header: Header for the value column.

        Returns:
            Self for method chaining.
        """
        rows = [[str(k), str(v)] for k, v in data.items()]
        return self.add_table(rows, headers=[key_header, value_header])

    # =========================================================================
    # IMAGE METHODS
    # =========================================================================

    def add_image(
        self,
        image_path: str,
        width: Optional[float] = None,
        height: Optional[float] = None,
        caption: Optional[str] = None,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    ) -> "DocumentBuilder":
        """
        Add an image.

        Args:
            image_path: Path to the image file.
            width: Image width in inches.
            height: Image height in inches.
            caption: Optional caption text.
            alignment: Alignment of the image.

        Returns:
            Self for method chaining.
        """
        self._doc.add_image(
            image_path=image_path,
            width=width,
            height=height,
            caption=caption,
            alignment=alignment,
        )
        return self

    # =========================================================================
    # SPECIAL CONTENT METHODS
    # =========================================================================

    def add_quote(
        self,
        text: str,
        author: Optional[str] = None,
    ) -> "DocumentBuilder":
        """
        Add a blockquote.

        Args:
            text: The quote text.
            author: Optional attribution.

        Returns:
            Self for method chaining.
        """
        self._doc.add_quote(text, author)
        return self

    def add_code_block(
        self,
        code: str,
        language: Optional[str] = None,
    ) -> "DocumentBuilder":
        """
        Add a code block.

        Args:
            code: The code text.
            language: Optional language identifier.

        Returns:
            Self for method chaining.
        """
        self._doc.add_code_block(code, language)
        return self

    def add_horizontal_line(self) -> "DocumentBuilder":
        """
        Add a horizontal line.

        Returns:
            Self for method chaining.
        """
        self._doc.add_horizontal_line()
        return self

    def add_page_break(self) -> "DocumentBuilder":
        """
        Add a page break.

        Returns:
            Self for method chaining.
        """
        self._doc.add_page_break()
        return self

    def add_table_of_contents(
        self,
        title: str = "Table of Contents",
        levels: int = 3,
    ) -> "DocumentBuilder":
        """
        Add a table of contents.

        Args:
            title: Title for the TOC section.
            levels: Number of heading levels to include.

        Returns:
            Self for method chaining.
        """
        self._doc.add_table_of_contents(title, levels)
        return self

    # =========================================================================
    # TEMPLATE METHODS
    # =========================================================================

    def replace_placeholders(self, replacements: dict[str, str]) -> "DocumentBuilder":
        """
        Replace {{PLACEHOLDER}} tokens in the document.

        Args:
            replacements: Dictionary of placeholder names to values.

        Returns:
            Self for method chaining.
        """
        self._doc.replace_placeholders(replacements)
        return self

    def with_template(self, template_name: str) -> "DocumentBuilder":
        """
        Apply a built-in template style.

        Args:
            template_name: Name of the template to apply.

        Returns:
            Self for method chaining.
        """
        from will.templates import get_template

        template_config = get_template(template_name)
        if template_config:
            # Apply template settings
            if "margins" in template_config:
                self.set_margins(preset=template_config["margins"])
            if "page_size" in template_config:
                self.set_page_size(preset=template_config["page_size"])
            # Additional template configuration can be applied here

        return self

    # =========================================================================
    # OUTPUT METHODS
    # =========================================================================

    def save(self, path: Union[str, Path]) -> "DocumentBuilder":
        """
        Save the document to a file.

        Args:
            path: Output file path.

        Returns:
            Self for method chaining.
        """
        self._doc.save(path)
        return self

    def to_bytes(self) -> bytes:
        """
        Get the document as bytes.

        Returns:
            Document content as bytes.
        """
        return self._doc.to_bytes()

    def get_info(self) -> dict[str, Any]:
        """
        Get document information and statistics.

        Returns:
            Dictionary containing document info.
        """
        return self._doc.get_info()

    # =========================================================================
    # CONTEXT MANAGER SUPPORT
    # =========================================================================

    def __enter__(self) -> "DocumentBuilder":
        """Enter context manager."""
        return self

    def __exit__(self, exc_type, exc_val, exc_tb) -> None:
        """Exit context manager."""
        pass


# =============================================================================
# CONVENIENCE FUNCTIONS
# =============================================================================

def create_document(
    title: Optional[str] = None,
    subtitle: Optional[str] = None,
    author: Optional[str] = None,
    template: Optional[str] = None,
) -> DocumentBuilder:
    """
    Create a new document with optional initial settings.

    Args:
        title: Document title.
        subtitle: Document subtitle.
        author: Document author.
        template: Template file path.

    Returns:
        A DocumentBuilder instance.
    """
    doc = WordDocument(template=template, title=title, author=author)
    builder = DocumentBuilder(doc)

    if title:
        builder.set_title(title, subtitle)
        builder.add_page_break()

    return builder


def quick_report(
    title: str,
    sections: list[dict[str, Any]],
    output_path: str,
    author: Optional[str] = None,
    include_toc: bool = False,
) -> DocumentBuilder:
    """
    Quickly create a report document.

    Args:
        title: Report title.
        sections: List of section dictionaries.
        output_path: Output file path.
        author: Optional author name.
        include_toc: Whether to include table of contents.

    Returns:
        The DocumentBuilder instance.

    Example:
        >>> quick_report(
        ...     title="Quarterly Report",
        ...     sections=[
        ...         {"title": "Overview", "text": "This quarter..."},
        ...         {"title": "Results", "bullets": ["Revenue up", "Costs down"]},
        ...     ],
        ...     output_path="report.docx"
        ... )
    """
    builder = create_document(title=title, author=author)

    if include_toc:
        builder.add_table_of_contents()

    for section in sections:
        if section.get("page_break"):
            builder.add_page_break()

        if section.get("title"):
            level = section.get("level", 1)
            builder.add_heading(section["title"], level=level)

        if section.get("text"):
            builder.add_paragraph(section["text"])

        if section.get("bullets"):
            builder.add_bullets(section["bullets"])

        if section.get("numbered"):
            builder.add_numbered_list(section["numbered"])

        if section.get("table"):
            table_data = section["table"]
            builder.add_table(
                data=table_data.get("data", []),
                headers=table_data.get("headers"),
            )

        if section.get("image"):
            img = section["image"]
            builder.add_image(
                image_path=img.get("path", img) if isinstance(img, dict) else img,
                caption=img.get("caption") if isinstance(img, dict) else None,
            )

    builder.save(output_path)
    return builder
